import re
from collections import Counter
from pathlib import Path
from typing import NamedTuple

from docx import Document as DocxDocument
from pypdf import PdfReader

STOPWORDS = {
    "about",
    "after",
    "again",
    "against",
    "because",
    "before",
    "between",
    "could",
    "their",
    "there",
    "these",
    "those",
    "through",
    "under",
    "which",
    "while",
    "with",
    "would",
    "study",
    "student",
    "system",
    "using",
    "from",
    "into",
    "that",
    "this",
    "have",
    "your",
    "will",
    "been",
    "also",
    "than",
    "them",
    "they",
    "over",
    "only",
}


class ExtractedPage(NamedTuple):
    page_number: int
    text: str


class ProcessedChunk(NamedTuple):
    text: str
    topic_label: str
    page_start: int | None
    page_end: int | None


class ProcessedDocument(NamedTuple):
    cleaned_text: str
    chunks: list[ProcessedChunk]
    topics: list[str]
    word_count: int


def extract_pages(path: Path) -> list[ExtractedPage]:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        reader = PdfReader(str(path))
        return [
            ExtractedPage(page_number=index + 1, text=page.extract_text() or "")
            for index, page in enumerate(reader.pages)
        ]
    if suffix == ".txt":
        raw_text = path.read_text(encoding="utf-8", errors="ignore")
        parts = raw_text.split("\f")
        return [ExtractedPage(index + 1, part) for index, part in enumerate(parts)]
    if suffix == ".docx":
        doc = DocxDocument(str(path))
        return [ExtractedPage(1, "\n".join(paragraph.text for paragraph in doc.paragraphs))]
    raise ValueError(f"Unsupported file type: {suffix}")


def extract_text(path: Path) -> str:
    return "\n\n".join(page.text for page in extract_pages(path))


def clean_text(raw_text: str) -> str:
    text = raw_text.replace("\x00", " ")
    text = re.sub(r"[\r\f\v]+", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def chunk_text(text: str, words_per_chunk: int = 420, overlap: int = 70) -> list[str]:
    words = text.split()
    if not words:
        return []

    chunks: list[str] = []
    step = max(words_per_chunk - overlap, 1)
    for start in range(0, len(words), step):
        chunk_words = words[start : start + words_per_chunk]
        if not chunk_words:
            continue
        chunks.append(" ".join(chunk_words))
        if start + words_per_chunk >= len(words):
            break
    return chunks


def _topic_from_text(text: str) -> str | None:
    heading_match = re.search(r"(?:^|\s)Page\s+\d+\s*:\s*([A-Z][A-Za-z0-9, \-]+)", text)
    if heading_match:
        return heading_match.group(1).strip()[:120]

    keyword_match = re.search(r"Keywords:\s*([^.\n]+)", text, flags=re.IGNORECASE)
    if keyword_match:
        first_keyword = keyword_match.group(1).split(",")[0].strip()
        if first_keyword:
            return first_keyword.title()[:120]
    return None


def chunk_pages(
    pages: list[ExtractedPage],
    words_per_chunk: int = 260,
    overlap: int = 45,
) -> list[ProcessedChunk]:
    processed_chunks: list[ProcessedChunk] = []

    for page in pages:
        cleaned_page = clean_text(page.text)
        if not cleaned_page:
            continue

        topic_label = _topic_from_text(cleaned_page)
        page_chunks = chunk_text(cleaned_page, words_per_chunk=words_per_chunk, overlap=overlap)
        for page_chunk in page_chunks:
            processed_chunks.append(
                ProcessedChunk(
                    text=page_chunk,
                    topic_label=topic_label or "General",
                    page_start=page.page_number,
                    page_end=page.page_number,
                )
            )

    return processed_chunks


def _rank_keywords(text: str, limit: int = 8) -> list[str]:
    words = re.findall(r"[A-Za-z][A-Za-z\-]{3,}", text.lower())
    counts = Counter(word for word in words if word not in STOPWORDS)
    return [word.title() for word, _ in counts.most_common(limit)]


def extract_topics(text: str, limit: int = 5) -> list[str]:
    heading_topics = [
        match.strip()[:120]
        for match in re.findall(r"(?:^|\s)Page\s+\d+\s*:\s*([A-Z][A-Za-z0-9, \-]+)", text)
    ]
    keyword_topics = [
        keyword.strip().title()[:120]
        for keyword_block in re.findall(r"Keywords:\s*([^.\n]+)", text, flags=re.IGNORECASE)
        for keyword in keyword_block.split(",")
        if keyword.strip()
    ]
    topics = list(dict.fromkeys([*heading_topics, *keyword_topics]))[:limit]
    if len(topics) < limit:
        topics.extend(topic for topic in _rank_keywords(text, limit=limit) if topic not in topics)
    return topics or ["General"]


def choose_chunk_topic(chunk: str, global_topics: list[str]) -> str:
    lowered = chunk.lower()
    for topic in global_topics:
        if topic.lower() in lowered:
            return topic
    return global_topics[0] if global_topics else "General"


def process_document(path: Path) -> ProcessedDocument:
    pages = extract_pages(path)
    cleaned_pages = [clean_text(page.text) for page in pages if clean_text(page.text)]
    cleaned = "\n\n".join(cleaned_pages).strip()
    word_count = len(cleaned.split())
    topics = extract_topics(cleaned)
    chunks = []
    for chunk in chunk_pages(pages):
        topic = chunk.topic_label
        if topic == "General":
            topic = choose_chunk_topic(chunk.text, topics)
        chunks.append(
            ProcessedChunk(
                text=chunk.text,
                topic_label=topic,
                page_start=chunk.page_start,
                page_end=chunk.page_end,
            )
        )
    return ProcessedDocument(
        cleaned_text=cleaned,
        chunks=chunks,
        topics=topics,
        word_count=word_count,
    )
